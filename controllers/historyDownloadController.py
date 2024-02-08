from flask import Blueprint, send_file
from models.history import History, db

# for save_pdf
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

#for save_lontar_1line and save_lontar_2col
from docx import Document
from docx.shared import Inches, Pt
import os


# Create a Blueprint object for the controller
history_download_controller = Blueprint('history_download_controller', __name__)


# save to book format: .pdf
@history_download_controller.route('/save_book/<int:history_id>', methods=['GET'])
def save_book(history_id):
    # get data history by id    
    history = History.query.get(history_id)
    title = history.judul

    # Get the current time
    now = datetime.now()
    # Format the time
    current_date = now.strftime("%d-%m-%Y")

    # Buat dokumen PDF
    doc = SimpleDocTemplate(title+" "+current_date+".pdf", pagesize=letter)
    elements = []
    # Mendefinisikan jenis font "Bali Simbar"
    pdfmetrics.registerFont(TTFont("BaliSimbar", "static/font/balisdn.ttf"))

    # Buat gaya teks
    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.leading = 42
    style.fontSize = 18
    style.fontName = "BaliSimbar"

    # Tambahkan teks ke dokumen
    text = history.teks_hasil
    paragraph = Paragraph(text, style)
    elements.append(paragraph)

    # Build the PDF document
    doc.build(elements)

    return send_file(title+" "+current_date+".pdf", as_attachment=True)

# save to lontar 1 line: .docx
@history_download_controller.route('/save_lontar_1line/<int:history_id>', methods=['GET'])
def save_lontar_1line(history_id):
   #get data history by id
    data = History.query.get(history_id)
    judul = data.judul

    # Get the current time
    now = datetime.now()
    # Format the time
    current_date = now.strftime("%d-%m-%Y")

    if not judul:
        return "No content to export to Lontar Format"
    else:
        # find data from database history with judul
        data = History.query.filter_by(judul=judul).first()
    
    # get teks_hasil frm data
    text_result = data.teks_hasil

    # Path ke file template Word
    template_path = os.path.join(os.getcwd(), 'TemplateLontarWatermark.docx')

    # Buat dokumen Word baru dengan menggunakan template
    doc = Document(template_path)

    # Set ukuran kertas
    section = doc.sections[0]

    # Set margin halaman menjadi 0 di seluruh sisi
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)

    # Gaya teks untuk konten
    content_style = doc.styles.add_style('CustomBodyText', 1)  # Create a new custom style
    # content_style.font.size = Pt(20)
    content_style.font.size = Pt(32)
    # content_style.font.name = 'poppins'  # Ganti dengan jenis font yang diinginkan
    content_style.font.name = 'bali simbar'  # Ganti dengan jenis font yang diinginkan

    # Konten teks
    paragraphs = text_result.split('\n')  # Split teks menjadi paragraf

    # Jika sudah ada 4 baris teks, tambahkan halaman baru
    if len(doc.paragraphs) % 4 == 0:
        doc.add_page_break()

    # paragraph = doc.add_paragraph(paragraph_text.strip())
    paragraph = doc.add_paragraph(paragraphs)
    paragraph.style = content_style

    # Simpan dokumen ke file
    doc.save(judul+':'+current_date+'.docx')

    # Kirim file ke pengguna
    return send_file(judul+':'+current_date+'.docx', as_attachment=True)

# save to lontar 2 column: .docx
@history_download_controller.route('/save_lontar_2col/<int:history_id>', methods=['GET'])
def save_lontar_2col(history_id):
    #get data history by id
    data = History.query.get(history_id)
    judul = data.judul

    # Get the current time
    now = datetime.now()
    # Format the time
    current_date = now.strftime("%d-%m-%Y")

    if not judul:
        return "No content to export to Lontar Format"
    else:
        # find data from database history with judul
        data = History.query.filter_by(judul=judul).first()
    
    # get teks_hasil frm data
    text_result = data.teks_hasil

    # Path ke file template Word
    # template_path = os.path.join(os.getcwd(), 'TemplateLontarWatermark.docx')
    template_path = os.path.join(os.getcwd(), 'TemplateLontarTwoCol1.docx')

    # Buat dokumen Word baru dengan menggunakan template
    doc = Document(template_path)

    # Set ukuran kertas
    section = doc.sections[0]

    # Set margin halaman menjadi 0 di seluruh sisi
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)

    # Gaya teks untuk konten
    content_style = doc.styles.add_style('CustomBodyText', 1)  # Create a new custom style
    # content_style.font.size = Pt(20)
    content_style.font.size = Pt(29)
    # content_style.font.name = 'poppins'  # Ganti dengan jenis font yang diinginkan
    content_style.font.name = 'bali simbar'  # Ganti dengan jenis font yang diinginkan
    # add line spacing 1.15
    content_style.paragraph_format.line_spacing = 1.15
    

    # Konten teks
    paragraphs = text_result.split('\n')  # Split teks menjadi paragraf
    for paragraph_text in paragraphs:
        
        # pecah bentuk array
        arrayTeks = []
        # Setiap 60 karakter, split teks dan masukkan ke dalam arrayTeks
        for i in range(0, len(paragraph_text), 60):
            arrayTeks.append(paragraph_text[i:i + 60])

        # for testing
        # print('0: '+arrayTeks[0])
        # print('1: '+arrayTeks[1])
        # print('2: '+arrayTeks[2])
        # print('3: '+arrayTeks[3])
        # print('4: '+arrayTeks[4])
        # print('5: '+arrayTeks[5])
        # print('6: '+arrayTeks[6])
        # print('7: '+arrayTeks[7])
        # print(len(arrayTeks))

        # store array to new array
        newArray = []
        # append array
        newArray.append(arrayTeks[0])
        newArray.append(arrayTeks[2])
        newArray.append(arrayTeks[4])
        newArray.append(arrayTeks[6])
        newArray.append(arrayTeks[1])
        newArray.append(arrayTeks[3])
        newArray.append(arrayTeks[5])
        newArray.append(arrayTeks[7])

        # add space at the end of every array items
        for i in range(len(newArray)):
            newArray[i] = newArray[i] + ' '

        # paragraph = doc.add_paragraph(paragraph_text.strip())
        paragraph = doc.add_paragraph(newArray)
        paragraph.style = content_style
            

    # Simpan dokumen ke file
    doc.save(judul+':'+current_date+'.docx')

    # Kirim file ke pengguna
    return send_file(judul+':'+current_date+'.docx', as_attachment=True)