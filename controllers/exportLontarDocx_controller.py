from flask import Blueprint, send_file, session
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
import os

# Import database models
from models.history import History, db

export_lontar_docx_blueprint = Blueprint('test_lontar', __name__)

@export_lontar_docx_blueprint.route('/test_lontar', methods=['GET'])
def test():
    # Retrieve the translated text from the session
    judul = session.get('title', '')

    # Get the current time
    now = datetime.now()
    # Format the time
    current_date = now.strftime("%d-%m-%Y")

    if not judul:
        return "No content to export to Lontar Format"
    else:
        # find data from database history with judul
        data = History.query.filter_by(judul=judul).first()
    
    # get teks_hasil frm data
    text_result = data.teks_hasil

    # Path ke file template Word
    # template_path = os.path.join(os.getcwd(), 'TemplateLontarWatermark.docx')
    template_path = os.path.join(os.getcwd(), 'TemplateLontarTwoCol1.docx')

    # Buat dokumen Word baru dengan menggunakan template
    doc = Document(template_path)

    # Set ukuran kertas
    section = doc.sections[0]

    # Set margin halaman menjadi 0 di seluruh sisi
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)

    # Gaya teks untuk konten
    content_style = doc.styles.add_style('CustomBodyText', 1)  # Create a new custom style
    # content_style.font.size = Pt(20)
    content_style.font.size = Pt(29)
    # content_style.font.name = 'poppins'  # Ganti dengan jenis font yang diinginkan
    content_style.font.name = 'bali simbar'  # Ganti dengan jenis font yang diinginkan
    # add line spacing 1.15
    content_style.paragraph_format.line_spacing = 1.15

    # Konten teks
    paragraphs = text_result.split('\n')  # Split teks menjadi paragraf
    for paragraph_text in paragraphs:
        # setiap 72 karakter, berikan enter
        # entertest = []    
        # entertest = '________'.join(paragraph_text[i:i+72] for i in range(0, len(paragraph_text), 72))
        # paragraph_text = entertest

        # pecah bentuk array
        arrayTeks = []
        # Setiap 60 karakter, split teks dan masukkan ke dalam arrayTeks
        for i in range(0, len(paragraph_text), 60):
            arrayTeks.append(paragraph_text[i:i + 60])

        # print('0: '+arrayTeks[0])
        # print('1: '+arrayTeks[1])
        # print('2: '+arrayTeks[2])
        # print('3: '+arrayTeks[3])
        # print('4: '+arrayTeks[4])
        # print('5: '+arrayTeks[5])
        # print('6: '+arrayTeks[6])
        # print('7: '+arrayTeks[7])

        # print(len(arrayTeks))

        newArray = []

        newArray.append(arrayTeks[0])
        newArray.append(arrayTeks[2])
        newArray.append(arrayTeks[4])
        newArray.append(arrayTeks[6])
        newArray.append(arrayTeks[1])
        newArray.append(arrayTeks[3])
        newArray.append(arrayTeks[5])
        newArray.append(arrayTeks[7])

        # print(newArray)
        # add space at the end of every array items
        for i in range(len(newArray)):
            newArray[i] = newArray[i] + ' '

        # paragraph = doc.add_paragraph(paragraph_text.strip())
        paragraph = doc.add_paragraph(newArray)
        paragraph.style = content_style

    # Simpan dokumen ke file
    doc.save(judul+':'+current_date+'.docx')

    # Kirim file ke pengguna
    return send_file(judul+':'+current_date+'.docx', as_attachment=True)
