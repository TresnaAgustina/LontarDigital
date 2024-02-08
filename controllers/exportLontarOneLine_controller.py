from flask import Blueprint, send_file, session
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
import os

# Import database models
from models.history import History, db

export_lontar_one_line_blueprint = Blueprint('lontar_one_line', __name__)

@export_lontar_one_line_blueprint.route('/one_line', methods=['GET'])
def index():
    # Retrieve the translated text from the session
    judul = session.get('title', '')

    # Get the current time
    now = datetime.now()
    # Format the time
    current_date = now.strftime("%d-%m-%Y")

    if not judul:
        return "No content to export to Lontar Format"
    else:
        # find data from database history with judul
        data = History.query.filter_by(judul=judul).first()
    
    # get teks_hasil frm data
    text_result = data.teks_hasil

    # Path ke file template Word
    template_path = os.path.join(os.getcwd(), 'TemplateLontarWatermark.docx')

    # Buat dokumen Word baru dengan menggunakan template
    doc = Document(template_path)

    # Set ukuran kertas
    section = doc.sections[0]

    # Set margin halaman menjadi 0 di seluruh sisi
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)

    # Gaya teks untuk konten
    content_style = doc.styles.add_style('CustomBodyText', 1)  # Create a new custom style
    # content_style.font.size = Pt(20)
    content_style.font.size = Pt(32)
    # content_style.font.name = 'poppins'  # Ganti dengan jenis font yang diinginkan
    content_style.font.name = 'bali simbar'  # Ganti dengan jenis font yang diinginkan

    # Konten teks
    paragraphs = text_result.split('\n')  # Split teks menjadi paragraf

    # Jika sudah ada 4 baris teks, tambahkan halaman baru
    if len(doc.paragraphs) % 4 == 0:
        doc.add_page_break()

    # paragraph = doc.add_paragraph(paragraph_text.strip())
    paragraph = doc.add_paragraph(paragraphs)
    paragraph.style = content_style

    # Simpan dokumen ke file
    doc.save(judul+':'+current_date+'.docx')

    # Kirim file ke pengguna
    return send_file(judul+':'+current_date+'.docx', as_attachment=True)